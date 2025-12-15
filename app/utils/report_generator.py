# app/utils/report_generator.py

import io
from datetime import datetime

import streamlit as st
import pandas as pd
import docx  # python-docx


def _build_report_text(
    df: pd.DataFrame,
    baseline_metrics: dict,
    setfit_metrics: dict,
) -> str:
    n_rows = len(df)
    total_spend = df["Debit"].sum() if "Debit" in df.columns else None

    if "PredictedCategory" in df.columns:
        class_counts = df["PredictedCategory"].value_counts().to_dict()
    else:
        class_counts = {}

    lines: list[str] = []

    lines.append("Helixsense – Facility Expense NLP Classifier")
    lines.append("Executive Summary")
    lines.append("")
    lines.append(f"Dataset size: {n_rows} transactions.")
    if total_spend is not None:
        lines.append(f"Total spend covered: {total_spend:,.0f} (all years).")
    lines.append(f"Silver‑label class distribution: {class_counts}")
    lines.append("")

    lines.append("Business problem")
    lines.append(
        "Finance and operations teams need a consistent way to classify ERP "
        "facility‑expense lines into Services, Equipment, and Material using "
        "only the free‑text Remarks column, so that spend can be analysed by "
        "category without manual tagging."
    )
    lines.append("")

    lines.append("Solution overview")
    lines.append(
        "- Data loading, cleaning, and feature engineering of the `data.xlsx` export."
    )
    lines.append(
        "- Weak‑supervision framework that combines keyword rules, cost bands, "
        "a zero‑shot transformer, and action‑verb patterns to generate silver labels."
    )
    lines.append(
        "- Baseline TF‑IDF + Logistic Regression model as a fast, interpretable benchmark."
    )
    lines.append(
        "- SetFit sentence‑transformer model fine‑tuned on silver labels for semantic understanding."
    )
    lines.append(
        "- Streamlit app with EDA, model explainability, comparison dashboard, and project‑aware chatbot."
    )
    lines.append("")

    lines.append("Key model results (validation)")
    lines.append(
        f"- Baseline ({baseline_metrics['model_name']}): "
        f"F1‑Weighted={baseline_metrics['f1_weighted']:.3f}, "
        f"F1‑Macro={baseline_metrics['f1_macro']:.3f}, "
        f"estimated cost from mis‑classified value="
        f"{baseline_metrics['business_cost']:,.0f}."
    )
    lines.append(
        f"- SetFit ({setfit_metrics['model_name']}): "
        f"F1‑Weighted={setfit_metrics['f1_weighted']:.3f}, "
        f"F1‑Macro={setfit_metrics['f1_macro']:.3f}, "
        f"estimated cost from mis‑classified value="
        f"{setfit_metrics['business_cost']:,.0f}."
    )
    lines.append("")

    lines.append("Interpretation")
    if (
        setfit_metrics["f1_weighted"] >= baseline_metrics["f1_weighted"]
        and setfit_metrics["business_cost"] <= baseline_metrics["business_cost"]
    ):
        lines.append(
            "SetFit delivers higher F1‑Weighted and lowers the estimated value of "
            "mis‑classified transactions versus the TF‑IDF baseline, particularly "
            "on high‑value service lines."
        )
        lines.append(
            "This makes SetFit the recommended production model when a small amount "
            "of transformer latency is acceptable."
        )
    else:
        lines.append(
            "The TF‑IDF baseline currently matches or exceeds SetFit on F1‑Weighted "
            "or cost‑weighted error, so it remains the safer choice until the "
            "semantic model is further improved."
        )
    lines.append("")

    lines.append("Operational considerations")
    lines.append(
        "- High‑confidence predictions from the weak‑supervision ensemble can be "
        "auto‑approved, while low‑confidence cases are queued for manual review."
    )
    lines.append(
        "- Cost‑weighted metrics should continue to be tracked in production, with "
        "alerts if the estimated mis‑classification cost drifts above agreed "
        "thresholds."
    )
    lines.append(
        "- Periodic re‑training with fresh data and refined labeling rules will "
        "improve robustness as facility‑expense patterns evolve."
    )

    return "\n".join(lines)


def _build_report_docx(
    df: pd.DataFrame,
    baseline_metrics: dict,
    setfit_metrics: dict,
) -> bytes:
    """Create a Word report in memory using python-docx. [web:40][web:37]"""
    text = _build_report_text(df, baseline_metrics, setfit_metrics)

    doc = docx.Document()
    doc.add_heading("Helixsense – Facility Expense NLP Report", 0)
    doc.add_paragraph(
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    doc.add_paragraph("")

    for line in text.split("\n"):
        if not line.strip():
            doc.add_paragraph("")
        elif line.endswith("Summary") or line in [
            "Business problem",
            "Solution overview",
            "Key model results (validation)",
            "Interpretation",
            "Operational considerations",
        ]:
            doc.add_heading(line, level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def render_report_tab(
    df: pd.DataFrame,
    baseline_metrics: dict,
    setfit_metrics: dict,
) -> None:
    """
    Render the Report tab.

    Shows a KPI-style dashboard and provides a downloadable Word report that
    management can attach to internal documentation. [file:160]
    """

    # KPI cards
    total_spend = df["Debit"].sum() if "Debit" in df.columns else None
    n_rows = len(df)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Transactions", f"{n_rows}")
    with col2:
        if total_spend is not None:
            st.metric("Total spend", f"{total_spend:,.0f}")
        else:
            st.metric("Total spend", "N/A")
    with col3:
        st.metric("Baseline F1‑Weighted", f"{baseline_metrics['f1_weighted']:.3f}")
    with col4:
        st.metric("SetFit F1‑Weighted", f"{setfit_metrics['f1_weighted']:.3f}")

    st.markdown("---")

    # Text report preview
    st.subheader("Executive summary (preview)")
    text = _build_report_text(df, baseline_metrics, setfit_metrics)
    st.text_area("Report preview", text, height=360)

    # Download as Word document
    docx_bytes = _build_report_docx(df, baseline_metrics, setfit_metrics)
    st.download_button(
        "Download report (.docx)",
        data=docx_bytes,
        file_name="helixsense_facility_expense_nlp_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    st.markdown(
        """
This report can be attached to internal documentation or presentations to explain
the problem, dataset, modelling approach, and why the chosen model is preferred
when judged on both F1‑Weighted and cost‑weighted error. [file:160]
"""
    )
